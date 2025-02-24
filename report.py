from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT
import matplotlib.pyplot as plt
import numpy as np
from scipy.interpolate import make_interp_spline, BSpline, splrep, PchipInterpolator, UnivariateSpline

def create_full_page_plot_with_text(output_image_path, text_data, x, y, add_text=False ):
    """
    Создает график на всю страницу с текстом в правом верхнем углу.

    :param output_image_path: Путь для сохранения изображения графика.
    :param text_data: Словарь с данными для текстового поля.
    """
    x_lim = np.array([0, 3, 9, 27])
    y_max = np.array([0, 970, 1050, 1350])
    y_min = np.array([0, 870, 930, 1110])

    # Вычисление значений на более густой сетке
    x_smooth = np.linspace(min(x), max(x), 100)

    # Локальная интерполяция
    pchip = PchipInterpolator(x, y)
    y_smooth = pchip(x_smooth)

    # Вычисление значений на более густой сетке
    x_lim_smooth = np.linspace(min(x_lim), max(x_lim), 100)

    # Локальная интерполяция
    pchip = PchipInterpolator(x_lim, y_max)
    y_max_smooth = pchip(x_lim_smooth)

    # Локальная интерполяция
    pchip = PchipInterpolator(x_lim, y_min)
    y_min_smooth = pchip(x_lim_smooth)

    # Создание графика
    fig, ax = plt.subplots(figsize=(11.69, 8.27))  # Размер A4 (альбомная ориентация)
    ax.scatter(x, y, marker='o', linestyle='-', color='blue', label='Усилия')
    ax.scatter(x_lim, y_max, marker='o', linestyle='-', color='black')
    ax.scatter(x_lim, y_min, marker='o', linestyle='-', color='black')
    ax.plot(x_smooth, y_smooth, color='red', label='График')
    ax.plot(x_lim_smooth, y_max_smooth, linestyle='--', color='brown', label='Допуск по ТУ')
    ax.plot(x_lim_smooth, y_min_smooth, linestyle='--', color='brown')
    ax.set_title("График по точкам")
    ax.set_xlabel("Ход штока S (мм)")
    ax.set_ylabel("Усилие P (кгс)")

    # Добавление подписей над точками
    for i in range(len(x_lim)):
        ax.annotate(
            y_max[i],  # Текст подписи
            (x_lim[i], y_max[i]),  # Координаты точки
            textcoords="offset points",  # Относительное положение текста
            xytext=(0, 10),  # Смещение текста вверх на 10 пикселей
            ha='center'  # Горизонтальное выравнивание по центру
        )
        ax.annotate(
            y_min[i],  # Текст подписи
            (x_lim[i], y_min[i]),  # Координаты точки
            textcoords="offset points",  # Относительное положение текста
            xytext=(0, -20),  # Смещение текста вверх на 10 пикселей
            ha='center'  # Горизонтальное выравнивание по центру
        )

    ax.legend()
    # Увеличение количества делений на осях
    ax.locator_params(axis='x', nbins=40)  # 10 делений по оси X
    ax.locator_params(axis='y', nbins=20)  # 15 делений по оси Y

    # Включение сетки
    ax.grid(True, which='both', linestyle='--', linewidth=1)
    # Добавление текстового поля в нижнем правом углу
    text_content = (f"ГРАФИК\n"
                    f"снятия характеристик зависимости\n"
                    f"усилия по штоку от хода числа оборотов с\n"
                    f"гидродемпфера 8-1930-700 №{text_data['num']}\n"
                    f"Sкр(мм) = 3 P = {text_data['3mm']} кгс\n"
                    f"Sкр(мм) = 6 P = {text_data['6mm']} кгс\n"
                    f"Sкр(мм) = 9 P = {text_data['9mm']} кгс\n"
                    f"Sкр(мм) = 16 P = {text_data['16mm']} кгс\n"
                    f"Sкр(мм) = 27 P = {text_data['27mm']} кгс\n"
                    f"Исполнитель: {text_data['name']}\n"
                    f"Контр. мастер: {text_data['master']}")

    if add_text:
        # Параметры текстового поля
        ax.text(
            0.8, 0.05,  # Координаты: 95% от ширины и 5% от высоты (нижний правый угол)
            text_content,
            transform=ax.transAxes,  # Используем систему координат осей (0 до 1)
            fontsize=12,
            verticalalignment='bottom',  # Выравнивание по нижней границе
            horizontalalignment='center',  # Выравнивание по правой границе
            bbox=dict(
                boxstyle="Square, pad=1",  # Форма рамки (закругленная) и отступ
                edgecolor="black",  # Цвет границы
                facecolor="white",  # Цвет фона
                linewidth=1  # Толщина границы
            )
        )

    # Удаление лишних полей вокруг графика
    plt.subplots_adjust(left=0.1, right=0.95, top=0.95, bottom=0.1)  # Оставляем место для текста

    # Сохранение графика
    plt.savefig(output_image_path, format='png', dpi=300, bbox_inches='tight')  # Высокое разрешение
    plt.close()

def set_landscape_orientation(doc):
    """
    Устанавливает альбомную ориентацию страницы.

    :param doc: Объект Document (документ Word).
    """
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width  # Меняем размеры страницы

def insert_full_page_image(doc, image_path):
    """
    Вставляет изображение на всю страницу.

    :param doc: Объект Document (документ Word).
    :param image_path: Путь к изображению графика.
    """
    section = doc.sections[0]
    page_width = section.page_width - section.left_margin - section.right_margin  # Ширина страницы без полей
    page_height = section.page_height - section.top_margin - section.bottom_margin  # Высота страницы без полей

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(image_path, width=page_width, height=page_height)  # Растягиваем изображение на всю страницу

def generate_report(template_path, output_path, plot_image_path, text_data, x, y, add_text):
    """
    Генерирует отчет с графиком на всю страницу и текстом в правом верхнем углу.

    :param template_path: Путь к шаблонному файлу Word.
    :param output_path: Путь для сохранения нового файла Word.
    :param plot_image_path: Путь к изображению графика.
    :param text_data: Словарь с данными для текстового поля.
    """
    # Создание графика с текстовым полем
    create_full_page_plot_with_text(plot_image_path, text_data, x, y, add_text=add_text)

    # Загрузка шаблона
    doc = Document(template_path)

    # Установка альбомной ориентации
    set_landscape_orientation(doc)

    # Вставка изображения графика на всю страницу
    insert_full_page_image(doc, plot_image_path)

    # Сохранение нового документа
    doc.save(output_path)

def text_convert(x, y):
    text = list.copy(["(нет данных)"] * 6)
    for i in range(6):
        if len(y) - 1 < i:
            return text
        text[i] = y[i]
    return text

# Основная часть программы
if __name__ == "__main__":
    template_path = "template.docx"  # Шаблон документа Word
    output_path = "output_report_full_page_landscape.docx"  # Итоговый документ Word
    plot_image_path = "full_page_plot.png"  # Изображение графика


    # Данные для графика
    # x = np.array([0, 3, 6, 9, 16, 27])
    # y = np.array([0, 900, 950, 990, 1100, 1200])

    # Данные для графика
    x = np.array([0, 3, 6, 9])
    y = np.array([0, 900, 950, 990])

    text = text_convert(x, y)

    # Данные для текстового поля
    text_data = {
        "num": "5547740",
        "3mm": text[1],
        "6mm": text[2],
        "9mm": text[3],
        "16mm": text[4],
        "27mm": text[5],
        "name": "Демин ИД",
        "master": "Демин ДИ"
    }

    # Генерация отчета
    generate_report(template_path, output_path, plot_image_path, text_data, x, y, add_text=True)